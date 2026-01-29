"""
售前投标AI辅助系统 - Web界面
"""
import os
import sys
import json
import tempfile
from typing import Dict, Any
from datetime import datetime
import streamlit as st

# 添加src到Python路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from graphs.graph import main_graph
from utils.file.file import File
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import shutil


# 页面配置
st.set_page_config(
    page_title="售前投标AI辅助系统",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 自定义CSS
st.markdown("""
<style>
    .main-title {
        font-size: 2.8rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 1rem;
        font-weight: bold;
    }
    .subtitle {
        font-size: 1.2rem;
        color: #666;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.6rem;
        color: #2e7d32;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
        font-weight: bold;
    }
    .info-box {
        background-color: #e3f2fd;
        padding: 1.5rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
        border-left: 4px solid #2196F3;
    }
    .success-box {
        background-color: #c8e6c9;
        padding: 1.5rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
        border-left: 4px solid #4CAF50;
    }
    .warning-box {
        background-color: #fff3cd;
        padding: 1.5rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
        border-left: 4px solid #FFC107;
    }
    .feature-card {
        background-color: #f5f5f5;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)


def save_uploaded_file(uploaded_file) -> str:
    """保存上传的文件到临时目录"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{uploaded_file.name}") as tmp_file:
            tmp_file.write(uploaded_file.getbuffer())
            return tmp_file.name
    except Exception as e:
        st.error(f"文件保存失败: {str(e)}")
        return None


def display_checklist_result(checklist: Dict[str, Any], section_title: str, color_class: str = "info-box"):
    """显示检查清单结果"""
    st.markdown(f"### {section_title}")
    st.markdown(f'<div class="{color_class}">', unsafe_allow_html=True)
    
    # 如果是字符串，直接显示
    if isinstance(checklist, str):
        st.markdown(checklist)
    elif isinstance(checklist, dict):
        for key, value in checklist.items():
            if isinstance(value, list):
                st.markdown(f"**{key}:**")
                for item in value:
                    if isinstance(item, dict):
                        for k, v in item.items():
                            st.markdown(f"- {k}: {v}")
                    else:
                        st.markdown(f"- {item}")
            else:
                st.markdown(f"**{key}:** {value}")
    elif isinstance(checklist, list):
        for item in checklist:
            if isinstance(item, dict):
                for k, v in item.items():
                    st.markdown(f"**{k}:** {v}")
            else:
                st.markdown(f"- {item}")
    
    st.markdown("</div>", unsafe_allow_html=True)


def generate_docx_report(result: Dict[str, Any]) -> bytes:
    """
    生成docx格式的分析报告

    Args:
        result: 分析结果字典

    Returns:
        docx文件的字节数据
    """
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading('投标文件智能分析报告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 生成时间
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()

    # 废标项检查
    doc.add_heading('一、废标项检查', level=1)
    invalid_items = result.get("invalid_items_check", "")
    if invalid_items:
        add_content_to_docx(doc, invalid_items)

    # 商务得分检查
    doc.add_heading('二、商务得分检查', level=1)
    commercial_score = result.get("commercial_score_check", "")
    if commercial_score:
        add_content_to_docx(doc, commercial_score)

    # 技术方案检查
    doc.add_heading('三、技术方案检查', level=1)
    technical_plan = result.get("technical_plan_check", "")
    if technical_plan:
        add_content_to_docx(doc, technical_plan)

    # 指标应答检查
    doc.add_heading('四、指标应答检查', level=1)
    indicator_response = result.get("indicator_response_check", "")
    if indicator_response:
        add_content_to_docx(doc, indicator_response)

    # 技术得分检查
    doc.add_heading('五、技术得分检查', level=1)
    technical_score = result.get("technical_score_check", "")
    if technical_score:
        add_content_to_docx(doc, technical_score)

    # 文件结构检查
    doc.add_heading('六、文件结构检查', level=1)
    bid_structure = result.get("bid_structure_check", "")
    if bid_structure:
        add_content_to_docx(doc, bid_structure)

    # 修改建议汇总
    doc.add_heading('七、修改建议汇总', level=1)
    summary = result.get("final_modification_suggestions", "")
    if summary:
        add_content_to_docx(doc, summary)

    # 保存到字节流
    from io import BytesIO
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    return doc_stream.getvalue()


def generate_material_docx(content: str, title: str) -> bytes:
    """
    生成材料内容的docx文件

    Args:
        content: 材料内容
        title: 材料标题

    Returns:
        docx文件的字节数据
    """
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading(title, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 生成时间
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()

    # 添加内容
    add_content_to_docx(doc, content)

    # 保存到字节流
    from io import BytesIO
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    return doc_stream.getvalue()


def add_content_to_docx(doc: Document, content: str):
    """
    将内容添加到docx文档中

    Args:
        doc: docx文档对象
        content: 要添加的内容
    """
    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 判断是否为标题（以===或##开头）
        if line.startswith('===') or line.startswith('#'):
            level = 2
            if line.startswith('===') and line.count('=') > 5:
                level = 1
            elif line.startswith('###'):
                level = 3
            doc.add_heading(line.lstrip('= #'), level=level)
        # 判断是否为列表项（以数字或-开头）
        elif line[0].isdigit() or (line[0] == '-' and len(line) > 1 and line[1].isspace()):
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)


def generate_pdf_report(result: Dict[str, Any]) -> bytes:
    """
    生成PDF格式的分析报告

    Args:
        result: 分析结果字典

    Returns:
        PDF文件的字节数据
    """
    from io import BytesIO

    # 创建PDF文档
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18
    )

    # 获取样式
    styles = getSampleStyleSheet()

    # 自定义样式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='darkblue',
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='darkgreen',
        spaceAfter=12,
        spaceBefore=20,
        fontName='Helvetica-Bold'
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        fontName='Helvetica',
        leading=14
    )

    # 构建内容
    story = []

    # 标题
    story.append(Paragraph("投标文件智能分析报告", title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
    story.append(Spacer(1, 20))

    # 定义各部分内容
    sections = [
        ("一、废标项检查", "invalid_items_check"),
        ("二、商务得分检查", "commercial_score_check"),
        ("三、技术方案检查", "technical_plan_check"),
        ("四、指标应答检查", "indicator_response_check"),
        ("五、技术得分检查", "technical_score_check"),
        ("六、文件结构检查", "bid_structure_check"),
        ("七、修改建议汇总", "final_modification_suggestions")
    ]

    # 添加各部分内容
    for section_title, key in sections:
        content = result.get(key, "")
        if content:
            story.append(Paragraph(section_title, heading_style))
            story.append(Spacer(1, 6))

            # 处理内容
            if isinstance(content, str):
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        # 转义特殊字符
                        line = line.replace('&', '&amp;')
                        line = line.replace('<', '&lt;')
                        line = line.replace('>', '&gt;')

                        # 处理标题标记
                        if line.startswith('===') or line.startswith('###'):
                            # 这是一个小标题
                            heading_text = line.lstrip('= #').strip()
                            story.append(Paragraph(heading_text, heading_style))
                        elif line.startswith('-') or (len(line) > 0 and line[0].isdigit() and line[1] == '.'):
                            # 这是一个列表项
                            story.append(Paragraph(f"• {line.lstrip('-0123456789. ')}", normal_style))
                        else:
                            # 普通段落
                            story.append(Paragraph(line, normal_style))

            story.append(Spacer(1, 12))

    # 生成PDF
    doc.build(story)
    buffer.seek(0)

    return buffer.getvalue()


def validate_and_fix_workflow_type(input_data: dict) -> dict:
    """
    确保input_data中的workflow_type是有效的值
    """
    if 'workflow_type' in input_data:
        wt = input_data['workflow_type']
        # 如果是中文或者其他无效值，转换为正确的值
        if wt == '投标材料生成' or wt not in ['check', 'generate']:
            input_data['workflow_type'] = 'generate'
            st.warning(f"⚠️ 自动修正了workflow_type值: {wt} → generate")
    return input_data


def main():
    """主函数"""
    
    # 标题
    st.markdown('<h1 class="main-title">🤖 售前投标AI辅助系统</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">智能分析招标文件，辅助投标材料生成，提升投标成功率</p>', unsafe_allow_html=True)
    
    # 侧边栏
    with st.sidebar:
        st.markdown("## 📚 知识库管理")
        
        # 知识库路径配置
        kb_path = st.text_input(
            "知识库路径",
            value="assets/knowledge_base",
            help="输入本地知识库文件夹路径，用于材料生成时检索历史资料"
        )
        st.session_state['kb_path'] = kb_path
        
        # 知识库操作按钮
        col_kb1, col_kb2 = st.columns(2)
        with col_kb1:
            if st.button("🔄 刷新索引", key="refresh_kb", use_container_width=True):
                if os.path.exists(kb_path):
                    try:
                        from src.tools.knowledge_base_tool import KnowledgeBaseTool
                        kb_tool = KnowledgeBaseTool()
                        kb_tool.initialize(kb_path)
                        st.success("✅ 知识库索引刷新成功！")
                    except Exception as e:
                        st.error(f"❌ 刷新失败: {str(e)}")
                else:
                    st.warning(f"⚠️ 路径不存在: {kb_path}")
        
        with col_kb2:
            if st.button("📁 检查路径", key="check_kb_path", use_container_width=True):
                if os.path.exists(kb_path):
                    st.success(f"✅ 路径有效")
                    st.info(f"📄 文件数: {len(os.listdir(kb_path))}")
                else:
                    st.warning(f"⚠️ 路径不存在: {kb_path}")
        
        st.markdown("---")
        st.markdown("## ⚠️ 注意事项")
        st.markdown("""
        - 仅支持PDF、Word、PPT格式
        - 文件大小不超过100MB
        - 分析/生成过程可能需要几分钟
        """)
    
    # 主内容区
    
    # ====== 功能选择区域 ======
    st.markdown("---")
    
    col_func1, col_func2 = st.columns(2)
    
    with col_func1:
        if st.button(
            "📊 投标文件检查",
            key="btn_check_mode",
            type="primary",
            use_container_width=True
        ):
            st.session_state['workflow_type'] = 'check'
            st.session_state['mode_selected'] = True
            st.rerun()
    
    with col_func2:
        if st.button(
            "✍️ 投标材料生成",
            key="btn_generate_mode",
            type="primary",
            use_container_width=True
        ):
            st.session_state['workflow_type'] = 'generate'
            st.session_state['mode_selected'] = True
            st.rerun()
    
    # 初始化状态
    if 'mode_selected' not in st.session_state:
        st.session_state['mode_selected'] = False
    if 'workflow_type' not in st.session_state:
        st.session_state['workflow_type'] = 'check'
    
    # 如果用户选择了模式，显示对应的功能界面
    if st.session_state.get('mode_selected', False):
        workflow_type = st.session_state['workflow_type']
        
        if workflow_type == 'check':
            # ====== 投标文件检查模式 ======
            st.markdown("---")
            st.markdown('<h2 class="section-header">📊 投标文件检查</h2>', unsafe_allow_html=True)
            
            # 功能说明
            st.markdown("""
            <div class="info-box">
                <strong>功能说明：</strong><br>
                上传招标文件和投标文件，系统将从六个维度智能分析投标文件，检测潜在问题和改进建议。
                <ul>
                    <li>✅ 废标项检测：识别可能导致废标的风险点</li>
                    <li>✅ 商务得分检查：评估商务部分的得分情况</li>
                    <li>✅ 技术方案评估：分析技术方案的完整性和竞争力</li>
                    <li>✅ 指标应答验证：检查指标应答的准确性</li>
                    <li>✅ 技术得分点分析：识别技术得分点</li>
                    <li>✅ 文件结构检查：验证文件结构的规范性</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown('<h3>📄 招标文件</h3>', unsafe_allow_html=True)
                tender_file = st.file_uploader(
                    "上传招标文件",
                    type=['pdf', 'docx', 'doc', 'pptx', 'ppt'],
                    key="tender_file_check"
                )
                if tender_file:
                    st.success(f"已选择: {tender_file.name}")
            
            with col2:
                st.markdown('<h3>📝 投标文件</h3>', unsafe_allow_html=True)
                bid_file = st.file_uploader(
                    "上传投标文件",
                    type=['pdf', 'docx', 'doc', 'pptx', 'ppt'],
                    key="bid_file_check"
                )
                if bid_file:
                    st.success(f"已选择: {bid_file.name}")
            
            # 分析按钮
            st.markdown("---")
            
            col_btn1, col_btn2 = st.columns([1, 1])
            with col_btn1:
                analyze_button = st.button("🚀 开始分析", type="primary", use_container_width=True)
            with col_btn2:
                if st.button("🔄 重新选择功能", use_container_width=True):
                    st.session_state['mode_selected'] = False
                    st.rerun()
            
            if analyze_button:
                if not tender_file or not bid_file:
                    st.error("❌ 请先上传招标文件和投标文件！")
                    return
                
                # 保存文件
                with st.spinner("正在保存文件..."):
                    tender_path = save_uploaded_file(tender_file)
                    bid_path = save_uploaded_file(bid_file)
                    
                    if not tender_path or not bid_path:
                        st.error("文件保存失败！")
                        return
                
                # 准备输入
                try:
                    input_data = {
                        "tender_file": {
                            "url": tender_path,
                            "file_type": "document"
                        },
                        "bid_file": {
                            "url": bid_path,
                            "file_type": "document"
                        },
                        "workflow_type": "check"
                    }
                    
                    st.success("文件准备就绪，开始分析...")
                    
                    # 运行工作流
                    with st.spinner("正在进行六维分析，请稍候..."):
                        result = main_graph.invoke(input_data)
                    
                    # 显示结果
                    st.markdown('<h2 class="section-header">📋 分析结果</h2>', unsafe_allow_html=True)

                    # 废标项检测结果
                    if result.get("invalid_items_check"):
                        invalid_items = result["invalid_items_check"]
                        # 检查是否包含废标风险关键词
                        if "未发现废标项" in invalid_items or "无废标风险" in invalid_items or "恭喜" in invalid_items:
                            st.markdown('<div class="success-box">✅ 未发现废标项，恭喜！</div>', unsafe_allow_html=True)
                        else:
                            display_checklist_result(invalid_items, "❌ 废标项检测结果", "warning-box")
                    
                    # 商务得分检查结果
                    if result.get("commercial_score_check"):
                        display_checklist_result(result["commercial_score_check"], "💰 商务得分检查")
                    
                    # 技术方案评估结果
                    if result.get("technical_plan_check"):
                        display_checklist_result(result["technical_plan_check"], "🔧 技术方案评估")
                    
                    # 指标应答验证结果
                    if result.get("indicator_response_check"):
                        display_checklist_result(result["indicator_response_check"], "📊 指标应答验证")
                    
                    # 技术得分点分析结果
                    if result.get("technical_score_check"):
                        display_checklist_result(result["technical_score_check"], "🎯 技术得分点分析")
                    
                    # 文件结构检查结果
                    if result.get("bid_structure_check"):
                        display_checklist_result(result["bid_structure_check"], "📁 文件结构检查")
                    
                    # 修改建议汇总
                    st.markdown('<h2 class="section-header">💡 修改建议汇总</h2>', unsafe_allow_html=True)
                    if result.get("modification_summary"):
                        summary = result["modification_summary"]
                        display_checklist_result(summary, "💡 修改建议汇总", "info-box")
                    
                    # 下载选项
                    st.markdown("---")
                    st.markdown('<h2 class="section-header">📥 下载分析报告</h2>', unsafe_allow_html=True)

                    # 保存结果到session_state
                    st.session_state['analysis_result'] = result

                    col1, col2 = st.columns(2)

                    # Word报告下载
                    with col1:
                        st.download_button(
                            label="📄 下载Word报告",
                            data=generate_docx_report(result),
                            file_name=f"投标文件分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_word_btn",
                            use_container_width=True
                        )

                    # PDF报告下载
                    with col2:
                        st.download_button(
                            label="📕 下载PDF报告",
                            data=generate_pdf_report(result),
                            file_name=f"投标文件分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            key="download_pdf_btn",
                            use_container_width=True
                        )
                    
                except Exception as e:
                    st.error(f"分析过程出错: {str(e)}")
                    st.error(f"错误详情: {type(e).__name__}")
                    import traceback
                    st.error(traceback.format_exc())
        
        else:
            # ====== 投标材料生成模式 ======
            st.markdown("---")
            st.markdown('<h2 class="section-header">✍️ 投标材料生成</h2>', unsafe_allow_html=True)
            
            # 功能说明
            st.markdown("""
            <div class="info-box">
                <strong>功能说明：</strong><br>
                上传招标文件，系统将智能分析招标文件要求，结合知识库和互联网搜索，生成商务或技术材料。
                <ul>
                    <li>✅ 商务材料：生成公司资质、项目经验、服务承诺等商务相关内容</li>
                    <li>✅ 技术材料：生成技术方案、系统架构、实施方案等技术相关内容</li>
                    <li>✅ 知识库检索：从本地知识库检索历史资料</li>
                    <li>✅ 互联网搜索：搜索最新的行业信息和参考资料</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown('<h3>📄 招标文件</h3>', unsafe_allow_html=True)
            tender_file = st.file_uploader(
                "上传招标文件",
                type=['pdf', 'docx', 'doc', 'pptx', 'ppt'],
                key="tender_file_generate"
            )
            if tender_file:
                st.success(f"已选择: {tender_file.name}")
            
            st.markdown("---")
            
            # 材料生成选项
            col_m1, col_m2 = st.columns(2)
            
            with col_m1:
                material_type = st.selectbox(
                    "选择材料类型",
                    ["commercial", "technical"],
                    format_func=lambda x: "💼 商务材料" if x == "commercial" else "🔧 技术材料"
                )
            
            with col_m2:
                use_kb = st.checkbox("使用知识库", value=True, help="是否从本地知识库检索相关材料")
            
            # 生成要求输入
            st.markdown('<h3>📝 材料生成要求（选填）</h3>', unsafe_allow_html=True)
            generation_requirements = st.text_area(
                "输入具体的材料生成要求",
                height=150,
                placeholder="例如：生成一份商务材料，重点阐述公司资质、项目经验和服务优势...",
                help="请详细描述您需要生成的内容，系统将根据招标文件要求和知识库内容智能生成"
            )
            
            # 生成按钮
            st.markdown("---")
            
            col_btn1, col_btn2 = st.columns([1, 1])
            with col_btn1:
                generate_button = st.button("✍️ 开始生成材料", type="primary", use_container_width=True)
            with col_btn2:
                if st.button("🔄 重新选择功能", use_container_width=True):
                    st.session_state['mode_selected'] = False
                    st.rerun()
            
            if generate_button:
                if not tender_file:
                    st.error("❌ 请先上传招标文件！")
                    return
                
                if not generation_requirements:
                    st.warning("⚠️ 建议输入具体的生成要求以获得更好的结果")
                
                # 保存文件
                with st.spinner("正在保存文件..."):
                    tender_path = save_uploaded_file(tender_file)
                    
                    if not tender_path:
                        st.error("文件保存失败！")
                        return
                
                # 准备输入
                try:
                    input_data = {
                        "tender_file": {
                            "url": tender_path,
                            "file_type": "document"
                        },
                        "workflow_type": "generate",
                        "material_type": material_type,
                        "generation_requirements": generation_requirements,
                        "kb_path": kb_path if use_kb else "",
                        "use_kb": use_kb
                    }
                    
                    st.success("文件准备就绪，开始生成材料...")
                    
                    # 运行工作流
                    with st.spinner(f"正在生成{material_type}材料，请稍候..."):
                        result = main_graph.invoke(input_data)
                    
                    # 显示生成结果
                    st.markdown('<h2 class="section-header">📋 生成结果</h2>', unsafe_allow_html=True)
                    
                    material_type_name = "商务" if material_type == "commercial" else "技术"
                    
                    # 检查是否有生成的材料
                    has_commercial_material = result.get("commercial_material", "")
                    has_technical_material = result.get("technical_material", "")
                    
                    if has_commercial_material or has_technical_material:
                        # 商务材料
                        if has_commercial_material:
                            with st.expander("💼 商务材料", expanded=True):
                                st.markdown('<div class="info-box">', unsafe_allow_html=True)
                                st.markdown(has_commercial_material)
                                st.markdown("</div>", unsafe_allow_html=True)
                                
                                # 下载按钮
                                col_d1, col_d2 = st.columns(2)
                                with col_d1:
                                    st.download_button(
                                        label="📄 下载Word文档",
                                        data=generate_material_docx(has_commercial_material, "商务材料"),
                                        file_name=f"商务材料_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="download_commercial_docx",
                                        use_container_width=True
                                    )
                                with col_d2:
                                    st.download_button(
                                        label="📄 下载文本",
                                        data=has_commercial_material,
                                        file_name=f"商务材料_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                                        mime="text/plain",
                                        key="download_commercial_txt",
                                        use_container_width=True
                                    )
                        
                        # 技术材料
                        if has_technical_material:
                            with st.expander("🔧 技术材料", expanded=True):
                                st.markdown('<div class="info-box">', unsafe_allow_html=True)
                                st.markdown(has_technical_material)
                                st.markdown("</div>", unsafe_allow_html=True)
                                
                                # 下载按钮
                                col_d1, col_d2 = st.columns(2)
                                with col_d1:
                                    st.download_button(
                                        label="📄 下载Word文档",
                                        data=generate_material_docx(has_technical_material, "技术材料"),
                                        file_name=f"技术材料_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="download_technical_docx",
                                        use_container_width=True
                                    )
                                with col_d2:
                                    st.download_button(
                                        label="📄 下载文本",
                                        data=has_technical_material,
                                        file_name=f"技术材料_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                                        mime="text/plain",
                                        key="download_technical_txt",
                                        use_container_width=True
                                    )
                    else:
                        # 兼容旧版本，显示generated_material字段
                        if result.get("generated_material"):
                            generated_material = result["generated_material"]
                            st.markdown('<div class="info-box">', unsafe_allow_html=True)
                            st.markdown(generated_material)
                            st.markdown("</div>", unsafe_allow_html=True)
                            
                            # 下载按钮
                            col_d1, col_d2 = st.columns(2)
                            with col_d1:
                                st.download_button(
                                    label=f"📄 下载{material_type_name}材料Word文档",
                                    data=generate_material_docx(generated_material, f"{material_type_name}材料"),
                                    file_name=f"{material_type_name}材料_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"download_{material_type}_docx",
                                    use_container_width=True
                                )
                            with col_d2:
                                st.download_button(
                                    label=f"📄 下载{material_type_name}材料文本",
                                    data=generated_material,
                                    file_name=f"{material_type_name}材料_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                                    mime="text/plain",
                                    key=f"download_{material_type}_txt",
                                    use_container_width=True
                                )
                        else:
                            st.warning("⚠️ 未能生成材料，请检查输入和要求")
                
                    except Exception as e:
                        st.error(f"材料生成过程出错: {str(e)}")
                        st.error(f"错误详情: {type(e).__name__}")
                        import traceback
                        st.error(traceback.format_exc())
    
    # 页脚
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 1rem;">
        <p>🤖 售前投标AI辅助系统 | 基于LangGraph工作流引擎</p>
        <p>💡 智能分析，精准生成，提升投标成功率</p>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
